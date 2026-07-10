"""Fill the .docx template with the job-tailored CV content.

The generated Markdown CV (same section structure as the template) is parsed into
sections and written into the template at the <w:t> (text node) level, so the
template's exact formatting — fonts, sizes, spacing, margins, bold labels, and the
project-title hyperlinks — is preserved. No extra bold is added: markdown ** markers
are stripped to plain text, matching the example CV exactly.

Untouchable areas (photo, name, tagline, contact line, the "Software Developer -
Başlangıç Noktası" experience, and the whole Education block) are never located, so
they are left exactly as in the template.
"""
from __future__ import annotations

import copy
import io
import re

from docx import Document
from docx.oxml.ns import qn

from .config import CV_TEMPLATE_FILE

W_T = qn("w:t")
W_P = qn("w:p")
XML_SPACE = qn("xml:space")


# --------------------------------------------------------------------------- #
# Markdown parsing
# --------------------------------------------------------------------------- #

def _strip_md(text: str) -> str:
    """Remove markdown bold/italic markers and trailing <br>, return plain text.

    Strips ** (bold) and * (italic, e.g. ``*(Ongoing)*``) so no markers leak into the
    Google Doc, which must match the example CV exactly with no extra emphasis.
    """
    text = text.replace("<br>", "").replace("**", "").replace("*", "")
    return text.strip()


_TITLE_LINK_RE = re.compile(r"^\[(?P<disp>.+?)\]\((?P<url>[^)\s]+)\)(?P<rest>.*)$")


def _parse_title_link(text: str):
    """A project title '[Display](https://github.com/...)' -> (display_text, url).

    Returns (plain_title, "") when there is no link. Bold/italic markers are stripped
    first, so a bold-wrapped title like '**[Display](url)**' still parses — otherwise
    the leading '**' breaks the ^\\[ anchor, the URL is lost, and the slot silently
    keeps the template's original (wrong) hyperlink.
    """
    m = _TITLE_LINK_RE.match(_strip_md(text))
    if m:
        display = (m.group("disp") + m.group("rest")).strip()
        return _strip_md(display), m.group("url").strip()
    return _strip_md(text), ""


def parse_markdown_cv(md: str) -> dict:
    """Parse the generated CV markdown into structured, plain-text sections."""
    lines = md.splitlines()
    section = None
    summary: list[str] = []
    core_skills: list[str] = []
    intern_bullets: list[str] = []
    previous_exp: list[str] = []
    projects: list[dict] = []
    cur_entry_is_intern = False
    cur_project: dict | None = None

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("### "):
            section = stripped[4:].strip().upper()
            cur_project = None
            cur_entry_is_intern = False
            continue
        if not stripped:
            continue
        if section is None:
            # Ignore any preamble before the first "### " header. Without this, a stray
            # line before "### SUMMARY" reaches the section.startswith() check below and
            # crashes with: 'NoneType' object has no attribute 'startswith'.
            continue

        if section == "SUMMARY":
            summary.append(_strip_md(stripped))

        elif section == "CORE SKILLS":
            core_skills.append(_strip_md(stripped))

        elif section == "EXPERIENCE":
            if stripped.startswith("#### "):
                title = stripped[5:]
                cur_entry_is_intern = "intern" in title.lower()
            elif cur_entry_is_intern and stripped.startswith("*"):
                intern_bullets.append(_strip_md(stripped.lstrip("*").strip()))

        elif section == "SELECTED PROJECTS":
            if stripped.startswith("#### "):
                disp, url = _parse_title_link(stripped[5:])
                cur_project = {"title": disp, "title_url": url, "bullets": [], "tech": ""}
                projects.append(cur_project)
            elif cur_project is not None and stripped.lower().startswith("tech:"):
                cur_project["tech"] = _strip_md(stripped.split(":", 1)[1])
            elif cur_project is not None and stripped.startswith("*"):
                cur_project["bullets"].append(_strip_md(stripped.lstrip("*").strip()))

        elif section.startswith("PREVIOUS EXPERIENCE"):
            # The role title now lives in the heading, so every line here is description.
            previous_exp.append(_strip_md(stripped))

    return {
        "summary": " ".join(summary).strip(),
        "core_skills": core_skills,
        "intern_bullets": intern_bullets,
        "projects": projects,
        # The role title is in the heading now, so all collected lines are the description.
        "previous_exp_desc": " ".join(previous_exp).strip(),
    }


# --------------------------------------------------------------------------- #
# Low-level docx helpers (operate on <w:t> nodes to preserve all formatting)
# --------------------------------------------------------------------------- #

def _t_nodes(p) -> list:
    """All <w:t> text nodes inside a paragraph, in document order (incl. hyperlinks)."""
    return p._p.findall(".//" + W_T)


def _para_text(p) -> str:
    return "".join((t.text or "") for t in _t_nodes(p))


def _set_node(node, text: str) -> None:
    node.text = text
    node.set(XML_SPACE, "preserve")


def _replace_content(p, new_text: str, keep_prefix_nodes: int = 0) -> None:
    """Put new_text into the first content text node, clearing later content nodes.

    keep_prefix_nodes: number of leading <w:t> nodes to leave untouched (e.g. a bold
    label like 'Backend:' or a bullet '• ' or 'Tech:').
    """
    nodes = _t_nodes(p)
    if len(nodes) <= keep_prefix_nodes:
        return
    content_nodes = nodes[keep_prefix_nodes:]
    # Preserve a leading space if the label node didn't end with one.
    prefix_text = nodes[keep_prefix_nodes - 1].text if keep_prefix_nodes else ""
    lead = " " if (prefix_text and not prefix_text.endswith(" ")) else ""
    _set_node(content_nodes[0], lead + new_text)
    for n in content_nodes[1:]:
        _set_node(n, "")


def _set_bullet_text(p, text: str) -> None:
    """Set a bullet paragraph's text, keeping the '• ' glyph and node[0] formatting.

    Handles both template layouts: glyph+text in one <w:t> node, or '• ' and content
    split across two nodes. We collapse everything into the first node (bullets are
    plain non-bold, so this preserves the exact look).
    """
    nodes = _t_nodes(p)
    if not nodes:
        return
    _set_node(nodes[0], "• " + text)
    for n in nodes[1:]:
        _set_node(n, "")


def _set_hyperlink_target(p, url: str) -> None:
    """Point the paragraph's hyperlink (e.g. a project title) at `url`."""
    hl = p._p.find(qn("w:hyperlink"))
    if hl is None:
        return
    rid = hl.get(qn("r:id"))
    rels = p.part.rels
    if rid and rid in rels:
        rels[rid]._target = url


def _is_bullet(p) -> bool:
    return _para_text(p).lstrip().startswith("•")


def _is_tech(p) -> bool:
    return _para_text(p).strip().lower().startswith("tech:")


def _clone_after(p):
    """Deep-copy paragraph p and insert the copy right after it. Returns the new CT_P."""
    new_p = copy.deepcopy(p._p)
    p._p.addnext(new_p)
    return new_p


def _remove(p_elem) -> None:
    p_elem.getparent().remove(p_elem)


# --------------------------------------------------------------------------- #
# Section fillers
# --------------------------------------------------------------------------- #

def _find_header(paragraphs, prefix: str) -> int:
    for i, p in enumerate(paragraphs):
        if _para_text(p).strip().upper().startswith(prefix):
            return i
    raise ValueError(f"Header not found: {prefix}")


def _fill_bullets(bullet_paras: list, new_bullets: list[str]) -> None:
    """Match the number of bullet paragraphs to new_bullets, then set their text.

    bullet_paras: existing template bullet Paragraph objects (>=1). The bullet glyph
    ('• ') sits in the first text node, which we keep.
    """
    if not bullet_paras or not new_bullets:
        return

    from docx.text.paragraph import Paragraph

    parent = bullet_paras[0]._parent
    elems = [bp._p for bp in bullet_paras]

    # Grow: clone the last bullet element as needed.
    while len(elems) < len(new_bullets):
        last = elems[-1]
        new_el = copy.deepcopy(last)
        last.addnext(new_el)
        elems.append(new_el)
    # Shrink: remove extra bullet elements.
    while len(elems) > len(new_bullets):
        _remove(elems.pop())

    for el, text in zip(elems, new_bullets):
        _set_bullet_text(Paragraph(el, parent), text)


def _fill_core_skills(skill_paras: list, lines: list[str]) -> None:
    """Fill the 'Label: content' skill paragraphs, one per markdown line.

    The template ships a fixed set of slots, but the CV may legitimately carry a
    different number of skill lines (the Frontend line is dropped for a pure back-end
    posting), so grow/shrink the slots like the bullet lists do. The bold LABEL is taken
    from the markdown too: filling only the content would silently shift every line
    under the template's original label once a line is added or removed.
    """
    if not skill_paras or not lines:
        return

    from docx.text.paragraph import Paragraph

    parent = skill_paras[0]._parent
    elems = [p._p for p in skill_paras]

    while len(elems) < len(lines):
        new_el = copy.deepcopy(elems[-1])
        elems[-1].addnext(new_el)
        elems.append(new_el)
    while len(elems) > len(lines):
        _remove(elems.pop())

    for el, line in zip(elems, lines):
        p = Paragraph(el, parent)
        nodes = _t_nodes(p)
        if not nodes:
            continue
        label, sep, content = line.partition(":")
        if not sep:  # no label on this line: drop it all into the content node
            _replace_content(p, line.strip(), keep_prefix_nodes=0)
            continue
        _set_node(nodes[0], f"{label.strip()}:")
        _replace_content(p, content.strip(), keep_prefix_nodes=1)


def build_docx(markdown_cv: str) -> bytes:
    """Return a .docx (bytes) = template filled with the job-tailored CV content."""
    data = parse_markdown_cv(markdown_cv)
    doc = Document(str(CV_TEMPLATE_FILE))
    paras = doc.paragraphs

    # --- SUMMARY: the first non-empty paragraph after the SUMMARY header ---
    i = _find_header(paras, "SUMMARY")
    for p in paras[i + 1:]:
        if _para_text(p).strip():
            _replace_content(p, data["summary"], keep_prefix_nodes=0)
            break

    # --- CORE SKILLS: one 'Label: content' line per markdown skill line ---
    cs_start = _find_header(paras, "CORE SKILLS")
    exp_start = _find_header(paras, "EXPERIENCE")
    skill_paras = [p for p in paras[cs_start + 1: exp_start] if _para_text(p).strip()]
    _fill_core_skills(skill_paras, data["core_skills"])

    # Slots may have been added/removed above, so re-read the paragraph list.
    paras = doc.paragraphs
    exp_start = _find_header(paras, "EXPERIENCE")

    # --- EXPERIENCE: only the "Intern" entry's bullets are editable ---
    proj_start = _find_header(paras, "SELECTED PROJECTS")
    intern_bullets: list = []
    in_intern = False
    for p in paras[exp_start + 1: proj_start]:
        txt = _para_text(p).strip()
        if not txt:
            continue
        if not _is_bullet(p):  # an entry title line
            in_intern = "intern" in txt.lower()
        elif in_intern:
            intern_bullets.append(p)
    if intern_bullets:
        _fill_bullets(intern_bullets, data["intern_bullets"])

    # --- SELECTED PROJECTS: positional fill of each project slot ---
    edu_start = _find_header(paras, "EDUCATION")
    # Group region into blocks: each block starts at a non-bullet/non-tech title line.
    blocks: list[dict] = []
    cur: dict | None = None
    for p in paras[proj_start + 1: edu_start]:
        if not _para_text(p).strip():
            continue
        if _is_bullet(p):
            if cur:
                cur["bullets"].append(p)
        elif _is_tech(p):
            if cur:
                cur["tech"] = p
        else:  # title line
            cur = {"title": p, "bullets": [], "tech": None}
            blocks.append(cur)

    for slot, proj in zip(blocks, data["projects"]):
        _replace_content(slot["title"], proj["title"], keep_prefix_nodes=0)
        if proj.get("title_url"):
            _set_hyperlink_target(slot["title"], proj["title_url"])
        if slot["bullets"]:
            _fill_bullets(slot["bullets"], proj["bullets"])
        if slot["tech"] is not None and proj["tech"]:
            _replace_content(slot["tech"], proj["tech"], keep_prefix_nodes=1)

    # --- PREVIOUS EXPERIENCE: the role title now lives in the header line, so the
    # first non-empty paragraph after the header is the description itself. ---
    if data["previous_exp_desc"]:
        pe_start = _find_header(paras, "PREVIOUS EXPERIENCE")
        for p in paras[pe_start + 1:]:
            if _para_text(p).strip():
                _replace_content(p, data["previous_exp_desc"], keep_prefix_nodes=0)
                break

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
